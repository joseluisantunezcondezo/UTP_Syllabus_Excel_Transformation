import os
import re
import unicodedata
from docx import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

import time
from io import BytesIO
import zipfile  # <-- para generar el ZIP

import pandas as pd
import streamlit as st

# ======================================================
# ESTILOS GLOBALES
# ======================================================

def apply_global_styles():
    st.markdown(
        """
        <style>
        /* Reducir tamaño base de fuente para que todo se vea más compacto */
        html {
            font-size: 11px;  /* en lugar de 16px */
        }

        [data-testid="stAppViewContainer"] { background: #f3f4f6; }
        [data-testid="stSidebar"] { background: #f9fafb; }

        .utp-hero {
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            border-radius: 18px;
            padding: 1.8rem 2.4rem;
            color: #ffffff;
            margin-bottom: 1.8rem;
            box-shadow: 0 18px 40px rgba(76, 81, 191, 0.35);
            display: flex;
            align-items: center;
            gap: 1.0rem;
        }

        .utp-hero-icon {
            width: 3.1rem;
            height: 3.1rem;
            border-radius: 999px;
            background: rgba(255,255,255,0.18);
            display: flex;
            align-items: center;
            justify-content: center;
            font-size: 2.0rem;
        }
        .utp-hero-title {
            font-weight: 700;
            font-size: 1.8rem;
            margin-bottom: 0.15rem;
        }
        .utp-hero-sub {
            font-size: 0.92rem;
            opacity: 0.96;
        }

        /* Sidebar branding */
        .utp-sidebar-brand {
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            border-radius: 18px;
            padding: 1.0rem 1.1rem;
            color: #ffffff;
            box-shadow: 0 14px 32px rgba(76, 81, 191, 0.35);
            margin-bottom: 1.3rem;
        }
        .utp-sidebar-brand-title {
            font-weight: 700;
            font-size: 1.05rem;
            margin-bottom: 0.2rem;
            display: flex;
            align-items: center;
            gap: 0.4rem;
        }
        .utp-sidebar-brand-subtitle {
            font-size: 0.82rem;
            opacity: 0.92;
        }

        /* Cards generales */
        .utp-step-card {
            border-radius: 14px;
            border: 1px solid #e5e7eb;
            padding: 1.1rem 1.3rem 1.15rem 1.3rem;
            margin-bottom: 1.0rem;
            background: #ffffff;
            box-shadow: 0 10px 25px rgba(15,23,42,0.05);
        }

        /* Cabecera pasos con estado */
        .utp-step-row {
            display: flex;
            align-items: center;
            justify-content: space-between;
            margin-bottom: 0.7rem;
        }
        .utp-step-main {
            display: flex;
            align-items: center;
            gap: 0.55rem;
            font-size: 1.0rem;
            font-weight: 600;
            color: #111827;
        }
        .utp-step-number {
            display: inline-flex;
            align-items: center;
            justify-content: center;
            width: 26px;
            height: 26px;
            border-radius: 999px;
            background: #4f46e5;
            color: #ffffff;
            font-size: 0.9rem;
            font-weight: 600;
            box-shadow: 0 3px 8px rgba(79,70,229,0.45);
        }
        .utp-step-status {
            padding: 0.18rem 0.7rem;
            border-radius: 999px;
            font-size: 0.78rem;
            font-weight: 500;
            border: 1px solid transparent;
            white-space: nowrap;
        }
        .utp-step-status-ok {
            background-color: #dcfce7;
            color: #166534;
            border-color: #bbf7d0;
        }
        .utp-step-status-error {
            background-color: #fee2e2;
            color: #b91c1c;
            border-color: #fecaca;
        }

        /* Cabecera simple */
        .utp-step-header-simple {
            display: flex;
            align-items: center;
            gap: 0.55rem;
            font-size: 1.0rem;
            font-weight: 600;
            color: #111827;
            margin-bottom: 0.7rem;
        }
        .utp-step-header-simple .utp-step-number {
            width: 26px;
            height: 26px;
            border-radius: 999px;
            background: #4f46e5;
            color: #ffffff;
            display: inline-flex;
            align-items: center;
            justify-content: center;
            font-size: 0.9rem;
            font-weight: 600;
        }

        /* DataFrame (si en futuro se usa) */
        .stDataFrame {
            border-radius: 10px;
            border: 1px solid #e5e7eb;
        }

        /* Botones */
        .stButton>button {
            border-radius: 999px;
            font-weight: 600;
            padding: 0.6rem 1.3rem;
            border: none;
            transition: all 0.2s ease;
        }
        .stButton>button:hover {
            transform: translateY(-1px);
            box-shadow: 0 10px 25px rgba(79,70,229,0.45);
        }

        /* Barra de progreso tipo "task" (como UTP - Broken Link Checker) */
        .progress-bar-ui-task {
            margin-top: 0.5rem;
            padding: 0.75rem 1rem;
            border-radius: 12px;
            border: 1px solid #e5e7eb;
            background: #ffffff;
            box-shadow: 0 8px 20px rgba(15,23,42,0.04);
        }
        .progress-bar-ui-task-header {
            display: flex;
            align-items: center;
            justify-content: space-between;
            margin-bottom: 0.4rem;
            font-size: 0.86rem;
            color: #111827;
        }
        .progress-bar-ui-task-title {
            font-weight: 600;
        }
        .progress-bar-ui-task-percentage {
            font-weight: 600;
            color: #2563eb;
        }
        .progress-bar-ui-task-bar-track {
            width: 100%;
            height: 0.62rem;
            border-radius: 999px;
            background: #e5e7eb;
            overflow: hidden;
            margin-bottom: 0.4rem;
        }
        .progress-bar-ui-task-bar-fill {
            height: 100%;
            border-radius: 999px;
            background: linear-gradient(90deg, #3b82f6, #22c55e);
            width: 0%;
            transition: width 0.18s ease-out;
        }
        .progress-bar-ui-task-sub {
            font-size: 0.8rem;
            color: #4b5563;
        }
        </style>
        """,
        unsafe_allow_html=True,
    )


# ======================================================
# HERO – CABECERAS
# ======================================================

def render_hero():
    """Hero general (Home)."""
    st.markdown(
        """
        <div class="utp-hero">
            <div class="utp-hero-icon">📄</div>
            <div>
                <div class="utp-hero-title">
                    Plataforma UTP - Syllabus to Excel Transformation
                </div>
                <div class="utp-hero-sub">
                    Convierte múltiples archivos de silabos en formato Word a tablas estructuradas de Excel de forma automática.
                </div>
            </div>
        </div>
        """,
        unsafe_allow_html=True,
    )


def render_hero_syllabus():
    """Hero específico para el módulo Convert Syllabus to Excel."""
    st.markdown(
        """
        <div class="utp-hero">
            <div class="utp-hero-icon">📊</div>
            <div>
                <div class="utp-hero-title">
                    Syllabus to Excel Transformation
                </div>
                <div class="utp-hero-sub">
                    Convierte  múltiples archivos de silabos a tablas estructuradas de Excel.
                </div>
            </div>
        </div>
        """,
        unsafe_allow_html=True,
    )


def get_sidebar_header_html(title: str, subtitle: str, icon: str) -> str:
    return f"""
    <div class="utp-sidebar-brand">
        <div class="utp-sidebar-brand-title">
            <span>{icon}</span><span>{title}</span>
        </div>
        <div class="utp-sidebar-brand-subtitle">
            {subtitle}
        </div>
    </div>
    """


def render_step_header_html(step_label: str, title: str, ok: bool) -> str:
    status_text = "Listo" if ok else "Pendiente"
    status_class = "utp-step-status-ok" if ok else "utp-step-status-error"
    return f"""
    <div class="utp-step-row">
        <div class="utp-step-main">
            <span class="utp-step-number">{step_label}</span>
            <span>{title}</span>
        </div>
        <div class="utp-step-status {status_class}">{status_text}</div>
    </div>
    """


def render_simple_step_header(step_label: str, title: str):
    st.markdown(
        f"""
        <div class="utp-step-header-simple">
            <span class="utp-step-number">{step_label}</span>
            <span>{title}</span>
        </div>
        """,
        unsafe_allow_html=True,
    )

# ======================================================
# UI – HOME
# ======================================================

def render_home():
    """Home con descripción y funcionalidades de la plataforma."""
    # Hero superior existente
    render_hero()

    # Contenido tipo ficha debajo del hero
    st.markdown('<div class="utp-step-card">', unsafe_allow_html=True)

    st.markdown(
        """
### 🏠 Home

Plataforma UTP - Transformación de Sílabos a Excel es una herramienta inteligente desarrollada para automatizar la conversión de documentos académicos (sílabos) en formato Word a datos estructurados en Excel.  

Su objetivo principal es optimizar los procesos de gestión de carga de datos en la plataforma de Diseña +, reduciendo significativamente el tiempo de procesamiento manual y garantizando la consistencia en la estructuración de la información.

### ✨ Funcionalidades Principales

- **Carga Múltiple de Documentos Word.** Permite cargar uno o varios archivos Word (.docx) de manera simultánea.  
- **Extracción Inteligente de Datos.** La plataforma identifica y extrae automáticamente las siguientes secciones clave:
    - 📅 **Cronograma de Actividades** (Unidades de aprendizaje, distribución por semanas y sesiones, temas específicos por clase, actividades y evaluaciones programadas)
    - 📊 **Sistema de Evaluación** (Tipos de evaluación —parciales, trabajos, prácticas—, pesos porcentuales de cada componente, observaciones y criterios de calificación)
    - 🎓 **Información Académica** (Datos generales del curso —carrera, créditos, horas—, logros generales y específicos de aprendizaje, metodología de enseñanza, unidades de aprendizaje y temarios detallados)

Genera archivos Excel con múltiples hojas organizadas:

- Cronograma completo de actividades  
- Sistema de evaluación estructurado  
- Datos generales del curso  
- Logros y metodología  
- Unidades de aprendizaje detalladas  
        """,
        unsafe_allow_html=False,
    )

    st.markdown("</div>", unsafe_allow_html=True)

# ======================================================
# LÓGICA – CONVERT SYLLABUS TO EXCEL
# ======================================================

def formatear_nombre(nombre: str) -> str:
    palabras = nombre.replace("_", " ").split()
    return " ".join(p.capitalize() for p in palabras).upper()

def extraer_cronograma_doc(doc: Document, nombre_archivo: str):
    """
    Extracción ULTRA robusta del cronograma (hoja 'Cronograma').

    - Localiza primero la sección 'CRONOGRAMA DE ACTIVIDADES' en el documento
      (párrafos + tablas en orden real).
    - Dentro de esa sección:
        * Si una tabla tiene encabezado estándar
          ['Unidad de aprendizaje','Semana','Sesión','Tema','Actividades y evaluaciones'],
          se lee 1:1 pero corrigiendo semanas/sesiones cuando hay celdas combinadas.
        * Tablas sin encabezado estándar se procesan con heurísticas de contenido
          (detección de unidad, semana, sesión, tema, evaluaciones).
    - Al final se:
        * Deduplican filas.
        * Se rellenan hacia adelante Unidad/Título de unidad.
        * Se agrupan filas por (Unidad, Título, Semana, Sesión), concatenando
          trozos partidos (tema y actividades) en una sola fila por combinación.
    - Si no se encuentra la sección 10, cae a un fallback que busca la tabla
      estándar en todo el documento.
    """

    import unicodedata
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.table import Table, _Cell
    from docx.text.paragraph import Paragraph

    # -------------------- Identificación de curso --------------------
    nombre_base = os.path.splitext(os.path.basename(nombre_archivo))[0]
    partes = nombre_base.split("_")
    cod_catalogo = partes[0] if len(partes) >= 2 else ""
    nombre_curso = formatear_nombre("_".join(partes[1:])) if len(partes) >= 2 else ""

    columnas_esperadas = [
        "Unidad de aprendizaje",
        "Semana",
        "Sesión",
        "Tema",
        "Actividades y evaluaciones",
    ]

    # -------------------- Helpers genéricos --------------------
    def nfd(s: str) -> str:
        return unicodedata.normalize("NFD", s or "")

    def deacc(s: str) -> str:
        return "".join(ch for ch in nfd(s) if unicodedata.category(ch) != "Mn")

    def norm_space(s: str) -> str:
        s = (s or "").replace("\u00A0", " ")
        s = re.sub(r"[¤•·]+", " ", s)
        s = re.sub(r"\s+", " ", s)
        return s.strip()

    def only_digits(x) -> str:
        m = re.search(r"\d{1,2}", str(x or ""))
        return m.group(0) if m else ""

    def iter_block_items(parent):
        """
        Itera párrafos y tablas preservando el orden del documento.
        (igual patrón que en extraer_unidades_logros_doc)
        """
        if isinstance(parent, _Cell):
            parent_elm = parent._tc
        else:
            parent_elm = parent._element.body
        for child in parent_elm.iterchildren():
            if isinstance(child, CT_P):
                yield Paragraph(child, parent)
            elif isinstance(child, CT_Tbl):
                yield Table(child, parent)

    def table_to_rows(t: Table):
        return [[norm_space(c.text) for c in r.cells] for r in t.rows]

    def guess_mapping(rows):
        """
        Heurística para tablas sin encabezado estándar.
        Intenta inferir qué columna es Unidad / Semana / Sesión / Tema / Actividades.
        """
        ncols = max((len(r) for r in rows), default=0)
        metrics = []
        for j in range(ncols):
            col = [(r[j] if j < len(r) else "") for r in rows[: min(30, len(rows))]]
            num_pat = re.compile(r"^\s*\d{1,2}\s*$")
            nums = sum(1 for v in col if num_pat.match(v))
            num_any = sum(1 for v in col if re.search(r"\d{1,2}", v))
            unidad_hits = sum(1 for v in col if re.search(r"(?i)\bunidad\b", v))
            eval_tokens = sum(
                1
                for v in col
                if re.search(
                    r"(?i)pr[aá]ctica|portafolio|caso|laboratorio|avance|participaci[oó]n|"
                    r"exposici[oó]n|resoluci[oó]n|entrega|parcial|final|evaluaci[oó]n",
                    v,
                )
            )
            avg_len = sum(len(v) for v in col) / max(1, len(col))
            metrics.append(
                dict(
                    j=j,
                    nums=nums,
                    num_any=num_any,
                    unidad=unidad_hits,
                    evalt=eval_tokens,
                    avg_len=avg_len,
                )
            )

        # columna "semana" = donde predominan números cortos
        sorted_by_nums = sorted(
        metrics, key=lambda m: (m["nums"], m["num_any"]), reverse=True
        )
        semana = sorted_by_nums[0]["j"] if sorted_by_nums else None
        sesion = sorted_by_nums[1]["j"] if len(sorted_by_nums) > 1 else None

        rem = [m for m in metrics if m["j"] not in (semana, sesion)]
        unidad = None
        if rem:
            unidad = max(rem, key=lambda m: (m["unidad"], m["avg_len"]))["j"]
            rem = [m for m in rem if m["j"] != unidad]

        if rem:
            ayev = max(rem, key=lambda m: m["evalt"])["j"]
            rem2 = [m for m in rem if m["j"] != ayev]
            tema = rem2[0]["j"] if rem2 else None
        else:
            tema = None
            ayev = None

        return {
            "unidad": unidad,
            "semana": semana,
            "sesion": sesion,
            "tema": tema,
            "ayev": ayev,
        }

    def dedup(rows):
        seen, out = set(), []
        for d in rows:
            key = (
                d["Unidad"],
                d["Titulo_unidad"],
                d["Semana"],
                d["Sesión"],
                d["Tema"],
                d["Actividades y evaluaciones"],
            )
            if key in seen:
                continue
            seen.add(key)
            out.append(d)
        return out

    def aggregate_rows(rows):
        """
        Unifica filas con la misma (Unidad, Título, Semana, Sesión),
        concatenando Tema y Actividades. Esto resuelve tablas donde
        Word parte una misma fila en múltiples renglones.
        """
        aggregated = {}
        order = []

        def add_piece(prev, piece):
            piece = (piece or "").strip()
            if not piece:
                return prev
            if not prev:
                return piece
            # Evitar duplicar si ya está incluído literalmente
            if piece in prev:
                return prev
            return prev + " " + piece

        for d in rows:
            key = (
                d["Unidad"],
                d["Titulo_unidad"],
                str(d["Semana"]),
                str(d["Sesión"]),
            )
            if key not in aggregated:
                aggregated[key] = {
                    "Cod_Catalogo": d["Cod_Catalogo"],
                    "Nombre_curso": d["Nombre_curso"],
                    "Unidad": d["Unidad"],
                    "Titulo_unidad": d["Titulo_unidad"],
                    "Semana": d["Semana"],
                    "Sesión": d["Sesión"],
                    "Tema": d["Tema"],
                    "Actividades y evaluaciones": d["Actividades y evaluaciones"],
                }
                order.append(key)
            else:
                agg = aggregated[key]
                agg["Tema"] = add_piece(agg["Tema"], d["Tema"])
                agg["Actividades y evaluaciones"] = add_piece(
                    agg["Actividades y evaluaciones"],
                    d["Actividades y evaluaciones"],
                )
        return [aggregated[k] for k in order]

    # -------------------- 1) Localizar sección CRONOGRAMA --------------------
    rx_start = re.compile(
        r"^\s*(?:\d+\.\s*)?cronograma\s+de\s+actividades\s*$", re.I
    )
    rx_end = re.compile(
        r"^\s*(?:\d+\.\s*|bibliograf[íi]a|referencias|anexos|recursos|docente|"
        r"datos\s+del\s+docente|contenidos)\b",
        re.I,
    )

    in_sec = False
    sec_blocks = []
    for blk in iter_block_items(doc):
        probe = ""
        if isinstance(blk, Paragraph):
            probe = blk.text or ""
        else:
            try:
                if blk.rows and blk.rows[0].cells:
                    probe = " ".join(c.text for c in blk.rows[0].cells)
            except Exception:
                probe = ""
        if not in_sec and rx_start.search(norm_space(probe)):
            # Encontramos el inicio de "CRONOGRAMA DE ACTIVIDADES"
            in_sec = True
            continue
        if in_sec and isinstance(blk, Paragraph) and rx_end.search(
            norm_space(blk.text or "")
        ):
            # Fin de la sección
            break
        if in_sec:
            sec_blocks.append(blk)

    resultados: list[dict] = []

    # ==================== 2) Procesar SOLO dentro de la sección ====================
    if sec_blocks:
        last_unidad_num = ""
        last_titulo_unidad = ""
        carry_semana = ""
        carry_sesion = ""

        for blk in sec_blocks:
            # 2.a Párrafos que nombran unidades (ej: "Unidad 1. Fundamentos...")
            if isinstance(blk, Paragraph):
                txt = norm_space(getattr(blk, "text", ""))
                m = re.search(
                    r"(?i)\bunidad\s*(\d{1,2})\s*[:\-\.]?\s*(.*)", txt
                )
                if m:
                    last_unidad_num = m.group(1).strip()
                    if m.group(2).strip():
                        last_titulo_unidad = m.group(2).strip()
                continue

            # 2.b Tablas dentro de la sección
            rows = table_to_rows(blk)
            if not rows:
                continue

            encabezados = rows[0]

            # ---- 2.b.1 Tablas con encabezado estándar ----
            if all(col in encabezados for col in columnas_esperadas):
                idx = {
                    col: encabezados.index(col) for col in columnas_esperadas
                }
                last_semana = ""
                last_sesion = ""

                for row in rows[1:]:
                    if not any(row):
                        continue

                    texto_unidad = (
                        row[idx["Unidad de aprendizaje"]]
                        if idx["Unidad de aprendizaje"] < len(row)
                        else ""
                    )
                    match_u = re.match(r"Unidad\s*(\d+)\s*(.*)", texto_unidad)
                    unidad_num = match_u.group(1) if match_u else last_unidad_num
                    titulo_unidad = (
                        match_u.group(2).strip()
                        if (match_u and match_u.group(2))
                        else (texto_unidad or last_titulo_unidad)
                    )

                    if unidad_num:
                        last_unidad_num = unidad_num
                    if titulo_unidad:
                        last_titulo_unidad = titulo_unidad

                    raw_semana = (
                        row[idx["Semana"]] if idx["Semana"] < len(row) else ""
                    )
                    raw_sesion = (
                        row[idx["Sesión"]] if idx["Sesión"] < len(row) else ""
                    )

                    semana = only_digits(raw_semana) or last_semana
                    sesion = only_digits(raw_sesion) or last_sesion

                    if semana:
                        last_semana = semana
                    if sesion:
                        last_sesion = sesion

                    tema = (
                        row[idx["Tema"]] if idx["Tema"] < len(row) else ""
                    )
                    ayev = (
                        row[idx["Actividades y evaluaciones"]]
                        if idx["Actividades y evaluaciones"] < len(row)
                        else ""
                    )

                    tema_final = norm_space(tema)
                    ayev_final = norm_space(ayev)

                    # requerimos semana + (tema o actividades)
                    if not (semana and (tema_final or ayev_final)):
                        continue

                    resultados.append(
                        {
                            "Cod_Catalogo": cod_catalogo,
                            "Nombre_curso": nombre_curso,
                            "Unidad": (unidad_num or last_unidad_num).strip(),
                            "Titulo_unidad": norm_space(
                                titulo_unidad or last_titulo_unidad
                            ),
                            "Semana": semana,
                            "Sesión": sesion,
                            "Tema": tema_final,
                            "Actividades y evaluaciones": ayev_final,
                        }
                    )

                # seguir con siguientes tablas de la sección
                continue

            # ---- 2.b.2 Tablas SIN encabezado estándar -> heurística ----
            mapping = guess_mapping(rows)

            for r in rows:
                # Unidad
                txt_uni = (
                    r[mapping["unidad"]]
                    if mapping["unidad"] is not None
                    and mapping["unidad"] < len(r)
                    else ""
                )
                u_num, u_title = "", ""
                if txt_uni:
                    m = re.search(
                        r"(?i)\bunidad\b\s*(\d{1,2})\s*(.*)", txt_uni
                    )
                    if m:
                        u_num, u_title = m.group(1).strip(), m.group(2).strip()
                    else:
                        m2 = re.match(
                            r"^\s*(\d{1,2})\s*[.\-: ]\s*(.+)$", txt_uni
                        )
                        if m2:
                            u_num, u_title = (
                                m2.group(1).strip(),
                                m2.group(2).strip(),
                            )
                        else:
                            if len(txt_uni) >= 4:
                                u_title = txt_uni

                if u_num or u_title:
                    last_unidad_num = u_num or last_unidad_num
                    last_titulo_unidad = u_title or last_titulo_unidad

                # Semana / Sesión con "carry" por celdas combinadas
                sem = (
                    r[mapping["semana"]]
                    if mapping["semana"] is not None
                    and mapping["semana"] < len(r)
                    else ""
                )
                ses = (
                    r[mapping["sesion"]]
                    if mapping["sesion"] is not None
                    and mapping["sesion"] < len(r)
                    else ""
                )

                sem = only_digits(sem) or carry_semana
                ses = only_digits(ses) if ses != "" else carry_sesion
                if sem:
                    carry_semana = sem
                if ses:
                    carry_sesion = ses

                # Tema / Actividades
                tema = (
                    r[mapping["tema"]]
                    if mapping["tema"] is not None
                    and mapping["tema"] < len(r)
                    else ""
                )
                ayev = (
                    r[mapping["ayev"]]
                    if mapping["ayev"] is not None
                    and mapping["ayev"] < len(r)
                    else ""
                )

                # si no hay tema pero sí algo que parece evaluación -> forzamos etiqueta
                if not tema and re.search(
                    r"(?i)pr[aá]ctica|portafolio|caso|laboratorio|avance|"
                    r"participaci[oó]n|entrega|parcial|final|evaluaci[oó]n",
                    ayev,
                ):
                    tema = "Evaluación"

                semana_final = sem
                sesion_final = ses
                tema_final = norm_space(tema)
                ayev_final = norm_space(ayev)

                if not (semana_final and (tema_final or ayev_final)):
                    continue

                resultados.append(
                    {
                        "Cod_Catalogo": cod_catalogo,
                        "Nombre_curso": nombre_curso,
                        "Unidad": (last_unidad_num or "").strip(),
                        "Titulo_unidad": norm_space(last_titulo_unidad or ""),
                        "Semana": semana_final,
                        "Sesión": sesion_final,
                        "Tema": tema_final,
                        "Actividades y evaluaciones": ayev_final,
                    }
                )

        if resultados:
            # 1) deduplicar
            resultados = dedup(resultados)

            # 2) forward-fill de Unidad y Título de unidad
            fallback_u = ""
            fallback_t = ""
            for d in resultados:
                if d["Unidad"] and not fallback_u:
                    fallback_u = d["Unidad"]
                if d["Titulo_unidad"] and not fallback_t:
                    fallback_t = d["Titulo_unidad"]
                if fallback_u and fallback_t:
                    break
            last_u = fallback_u
            last_t = fallback_t
            for d in resultados:
                if d["Unidad"]:
                    last_u = d["Unidad"]
                else:
                    d["Unidad"] = last_u
                if d["Titulo_unidad"]:
                    last_t = d["Titulo_unidad"]
                else:
                    d["Titulo_unidad"] = last_t

            # 3) Agrupar filas por (Unidad, Título, Semana, Sesión) para
            #    evitar filas "cortadas" en tablas complejas.
            resultados = aggregate_rows(resultados)

            # 4) Ordenar por semana/sesión
            try:
                resultados.sort(
                    key=lambda d: (
                        int(only_digits(d["Semana"]) or 0),
                        int(only_digits(d["Sesión"]) or 0),
                    )
                )
            except Exception:
                pass

            return resultados

    # ==================== 3) Fallback: buscar tabla estándar en todo el doc ====================
    resultados_fallback = []
    for tabla in doc.tables:
        if not tabla.rows:
            continue
        encabezados = [cell.text.strip() for cell in tabla.rows[0].cells]
        if all(col in encabezados for col in columnas_esperadas):
            idx = {col: encabezados.index(col) for col in columnas_esperadas}
            last_semana = ""
            last_sesion = ""
            for row in tabla.rows[1:]:
                celdas = row.cells
                texto_unidad = celdas[idx["Unidad de aprendizaje"]].text.strip()
                match_u = re.match(r"Unidad\s*(\d+)\s*(.*)", texto_unidad)
                unidad_num = match_u.group(1) if match_u else ""
                titulo_unidad = (
                    match_u.group(2).strip()
                    if match_u and match_u.group(2)
                    else texto_unidad
                )

                raw_semana = celdas[idx["Semana"]].text.strip()
                raw_sesion = celdas[idx["Sesión"]].text.strip()
                semana = only_digits(raw_semana) or last_semana
                sesion = only_digits(raw_sesion) or last_sesion
                if semana:
                    last_semana = semana
                if sesion:
                    last_sesion = sesion

                tema = celdas[idx["Tema"]].text.strip()
                ayev = (
                    celdas[idx["Actividades y evaluaciones"]].text.strip()
                )

                tema_final = norm_space(tema)
                ayev_final = norm_space(ayev)
                if not (semana and (tema_final or ayev_final)):
                    continue

                resultados_fallback.append(
                    {
                        "Cod_Catalogo": cod_catalogo,
                        "Nombre_curso": nombre_curso,
                        "Unidad": unidad_num,
                        "Titulo_unidad": titulo_unidad,
                        "Semana": semana,
                        "Sesión": sesion,
                        "Tema": tema_final,
                        "Actividades y evaluaciones": ayev_final,
                    }
                )

    if resultados_fallback:
        resultados_fallback = dedup(resultados_fallback)

        # forward-fill mínimo de Unidad / Título
        fallback_u = ""
        fallback_t = ""
        for d in resultados_fallback:
            if d["Unidad"] and not fallback_u:
                fallback_u = d["Unidad"]
            if d["Titulo_unidad"] and not fallback_t:
                fallback_t = d["Titulo_unidad"]
            if fallback_u and fallback_t:
                break
        last_u = fallback_u
        last_t = fallback_t
        for d in resultados_fallback:
            if d["Unidad"]:
                last_u = d["Unidad"]
            else:
                d["Unidad"] = last_u
            if d["Titulo_unidad"]:
                last_t = d["Titulo_unidad"]
            else:
                d["Titulo_unidad"] = last_t

        resultados_fallback = aggregate_rows(resultados_fallback)

        try:
            resultados_fallback.sort(
                key=lambda d: (
                    int(only_digits(d["Semana"]) or 0),
                    int(only_digits(d["Sesión"]) or 0),
                )
            )
        except Exception:
            pass

        return resultados_fallback

    # ==================== 4) Nada encontrado ====================
    return []


def extraer_tabla_evaluacion_doc(doc: Document, nombre_archivo: str) -> pd.DataFrame:
    all_rows = []
    header_final = None

    for table in doc.tables:
        rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        if not rows or len(rows) < 2:
            continue

        header = rows[0]
        header_normalized = [h.lower() for h in header]

        if "tipo" in header_normalized and "observación" in header_normalized:
            all_rows.extend(rows[1:])
            header_final = header
            break

    if all_rows and header_final:
        df_eval = pd.DataFrame(all_rows, columns=header_final)
        df_eval.insert(0, "Archivo", os.path.basename(nombre_archivo))
        return df_eval
    else:
        return pd.DataFrame()


def limpiar_texto_generico(texto: str) -> str:
    """Normaliza espacios y elimina caracteres especiales."""
    texto = re.sub(r"[¤]+", "", texto)
    texto = re.sub(r"\t+", " ", texto)
    texto = re.sub(r"\s+", " ", texto)
    return texto.strip()


def limpiar_valor_generico(valor: str) -> str:
    """Elimina títulos residuales como 1.2., 1.3., etc."""
    return re.sub(r"\s*1\.\d+\.*", "", valor).strip()


def extraer_datos_generales_doc(doc: Document, nombre_archivo: str) -> pd.DataFrame:
    nombre_archivo_base = os.path.basename(nombre_archivo)

    texto_parrafos = [p.text for p in doc.paragraphs if p.text.strip() != ""]

    texto_tablas = []
    for tabla in doc.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                if celda.text.strip():
                    texto_tablas.append(celda.text.strip())

    texto_completo = "\n".join(texto_parrafos + texto_tablas)
    texto_completo = limpiar_texto_generico(texto_completo)

    carrera_patron = r"(?:Carrera\s*:?\s*)(.*?)(?=\s*Créditos|Enseñanza|Horas)"
    creditos_patron = r"(?:Créditos?\s*:?\s*)([\s\S]*?)(?=\s*(?:1\.3\.|Enseñanza|Horas\s*semanales))"
    ensenanza_patron = r"(?:Enseñanza\s*de\s*curso\s*:?\s*)([\s\S]*?)(?=\s*1\.4\.|Horas\s*semanales)"
    horas_patron = r"(?:Horas\s*semanales\s*:?\s*)(\d+)"

    carreras_match = re.search(carrera_patron, texto_completo, re.IGNORECASE | re.DOTALL)
    creditos_match = re.search(creditos_patron, texto_completo, re.IGNORECASE)
    ensenanza_match = re.search(ensenanza_patron, texto_completo, re.IGNORECASE)
    horas_match = re.search(horas_patron, texto_completo, re.IGNORECASE)

    carreras = ""
    if carreras_match:
        raw_carreras = carreras_match.group(1).strip()
        raw_carreras = re.sub(r"(?i)^Carrera\s*:?", "", raw_carreras)
        raw_carreras = re.sub(r"\s*1\.\d+\.*", "", raw_carreras)
        lista_carreras = [c.strip() for c in re.split(r"\n|\s{2,}", raw_carreras) if c.strip()]
        carreras = ", ".join(lista_carreras)

    creditos = limpiar_valor_generico(creditos_match.group(1)) if creditos_match else ""
    ensenanza = limpiar_valor_generico(ensenanza_match.group(1)) if ensenanza_match else ""
    horas = limpiar_valor_generico(horas_match.group(1)) if horas_match else ""

    fila = {
        "Archivo": nombre_archivo_base,
        "1.1.Carrera:": carreras,
        "1.2.Créditos:": creditos,
        "1.3.Enseñanza de curso:": ensenanza,
        "1.4.Horas semanales:": horas
    }

    return pd.DataFrame([fila])


def limpiar_texto_logro(texto: str) -> str:
    """Convierte tabulaciones en saltos de línea y elimina caracteres especiales."""
    texto = re.sub(r"[¤]+", "", texto)
    texto = re.sub(r"\t+", "\n", texto)
    return texto.strip()


def extraer_logro_metodologia_doc(doc: Document, nombre_archivo: str) -> pd.DataFrame:
    nombre_archivo_base = os.path.basename(nombre_archivo)

    texto_parrafos = [p.text for p in doc.paragraphs if p.text.strip()]

    texto_tablas = []
    for tabla in doc.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                if celda.text.strip():
                    texto_tablas.append(celda.text.strip())

    texto_completo = "\n".join(texto_parrafos + texto_tablas)
    texto_completo = limpiar_texto_logro(texto_completo)

    logro_patron = (
        r"LOGRO GENERAL DE APRENDIZAJE\s*(.*?)(?=\n\s*\d*\.*\s*UNIDADES Y LOGROS ESPECÍFICOS DE APRENDIZAJE)"
    )
    metodologia_patron = r"METODOLOGÍA\s*(.*?)(?=\n\s*\d*\.*\s*SISTEMA DE EVALUACIÓN)"

    logro_match = re.search(logro_patron, texto_completo, re.IGNORECASE | re.DOTALL)
    metodologia_match = re.search(metodologia_patron, texto_completo, re.IGNORECASE | re.DOTALL)

    def limpiar_seccion(texto_sec: str) -> str:
        texto_sec = texto_sec.strip()
        if not texto_sec:
            return ""
        return texto_sec if texto_sec.endswith(".") else texto_sec + "."

    logro_texto = limpiar_seccion(logro_match.group(1)) if logro_match else ""
    metodologia_texto = limpiar_seccion(metodologia_match.group(1)) if metodologia_match else ""

    filas = [{
        "Archivo": nombre_archivo_base,
        "4. LOGRO GENERAL DE APRENDIZAJE": logro_texto,
        "6. METODOLOGÍA": metodologia_texto
    }]

    return pd.DataFrame(filas)


def extraer_pesos_doc(doc: Document, nombre_archivo: str) -> pd.DataFrame:
    nombre_archivo_base = os.path.basename(nombre_archivo)

    texto = "\n".join([p.text for p in doc.paragraphs if p.text.strip() != ""])
    patrones = re.findall(r"\((\d+)%\)\s*([A-Za-z0-9]+)", texto)

    filas = []
    for peso, tipo in patrones:
        filas.append({
            "Archivo": nombre_archivo_base,
            "Tipo": tipo,
            "Peso": peso
        })

    return pd.DataFrame(filas)


def extraer_unidades_logros_doc(doc: Document, nombre_archivo: str) -> pd.DataFrame:
    import re
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.table import Table, _Cell
    from docx.text.paragraph import Paragraph

    nombre_archivo_base = os.path.basename(nombre_archivo)

    def norm_line(s: str) -> str:
        s = re.sub(r"[¤•·]", " ", s or "")
        s = s.replace("\u00A0", " ")
        s = re.sub(r"\s+", " ", s)
        return s.strip()

    def push_unique(lst, text):
        t = norm_line(text)
        if t and (not lst or lst[-1] != t):
            lst.append(t)

    def collapse_full_repeats(s: str) -> str:
        s = norm_line(s)
        if not s:
            return s
        m = re.match(r"^(?P<x>.+?)(?:\s+\1){1,}$", s, flags=re.DOTALL)
        return m.group("x").strip() if m else s

    def iter_block_items(parent):
        if isinstance(parent, _Cell):
            parent_elm = parent._tc
        else:
            parent_elm = parent._element.body
        for child in parent_elm.iterchildren():
            if isinstance(child, CT_P):
                yield Paragraph(child, parent)
            elif isinstance(child, CT_Tbl):
                yield Table(child, parent)

    lineas = []

    def append_line_once(seq, s):
        t = (s or "").strip()
        if t and (not seq or seq[-1] != t):
            seq.append(t)

    for block in iter_block_items(doc):
        if isinstance(block, Paragraph):
            append_line_once(lineas, block.text)
        else:
            for row in block.rows:
                seen_tc = set()
                for cell in row.cells:
                    tc_id = id(cell._tc)
                    if tc_id in seen_tc:
                        continue
                    seen_tc.add(tc_id)
                    for p in cell.paragraphs:
                        append_line_once(lineas, p.text)

    hdr = re.compile(
        r"^(?:\s*5\.\s*)?UNIDADES\s+Y\s+LOGROS\s+ESPEC[ÍI]FICOS\s+DE\s+APRENDIZAJE\s*$",
        re.IGNORECASE,
    )
    rx_end = re.compile(
        r"^\s*(?:\d+\.\s*)?(?:METODOLOG[ÍI]A|SISTEMA\s+DE\s+EVALUACI[ÓO]N|FUENTES\s+DE\s+INFORMACI[ÓO]N|"
        r"BIBLIOGRAF[ÍI]A|COMPETENCIAS|CRONOGRAMA(?:\s+DE\s+ACTIVIDADES)?)\s*$",
        re.IGNORECASE,
    )

    in_section, s5 = False, []
    for ln in lineas:
        if not in_section:
            if hdr.search(ln):
                in_section = True
            continue
        if rx_end.search(ln):
            break
        s5.append(ln)

    if not s5:
        return pd.DataFrame(
            columns=["Archivo", "Unidad", "Unidades_Aprendizaje", "Logro_Especifico", "Temario"]
        )

    rx_unidad_full = re.compile(r"^\s*Unidad\s*de\s*aprendizaje\s*(\d+)\s*:?\s*$", re.IGNORECASE)
    rx_logro_full = re.compile(r"^\s*Logros?\s*espec[íi]ficos?\s*de\s*aprendizaje\s*:?\s*$", re.IGNORECASE)
    rx_temario_full = re.compile(r"^\s*Temario\s*:?\s*$", re.IGNORECASE)

    rx_unidad_inline = re.compile(
        r"^\s*Unidad\s*de\s*aprendizaje\s*(\d+)\s*:?\s*(.+)$",
        re.IGNORECASE,
    )
    rx_logro_inline = re.compile(
        r"^\s*Logros?\s*espec[íi]ficos?\s*de\s*aprendizaje\s*:?\s*(.+)$",
        re.IGNORECASE,
    )
    rx_temario_inline = re.compile(r"^\s*Temario\s*:?\s*(.+)$", re.IGNORECASE)

    rx_cut_inside = rx_end

    def limpia_semana(txt: str) -> str:
        txt = re.sub(r"(?im)^\s*Semana\s*\d+(?:\s*,\s*\d+)*(?:\s*y\s*\d+)?\s*$", "", txt)
        txt = re.sub(r"Semana\s*\d+(?:\s*,\s*\d+)*(?:\s*y\s*\d+)?", "", txt, flags=re.IGNORECASE)
        return txt

    registros = []
    curr, estado = None, None  # 'titulo' | 'logro' | 'temario'

    def cerrar_unidad():
        if not curr:
            return
        titulo = limpia_semana("\n".join(curr["titulo"]).strip())
        ua = norm_line(titulo)

        logro_full = norm_line(" ".join(curr["logro"])) if curr["logro"] else ""
        tem_full = norm_line(" ".join(curr["temario"])) if curr["temario"] else ""

        logro_full = collapse_full_repeats(logro_full)
        tem_full = collapse_full_repeats(tem_full)

        mcut = rx_cut_inside.search(tem_full)
        if mcut:
            tem_full = norm_line(tem_full[:mcut.start()])

        registros.append({
            "Archivo": nombre_archivo_base,
            "Unidad": curr["Unidad"],
            "Unidades_Aprendizaje": ua,
            "Logro_Especifico": logro_full,
            "Temario": tem_full
        })

    for raw in s5:
        ln = raw.strip()
        if not ln:
            continue

        if estado in ("temario", "logro") and rx_end.search(ln):
            cerrar_unidad()
            curr, estado = None, None
            break

        m = rx_unidad_full.match(ln)
        if m:
            cerrar_unidad()
            curr = {"Unidad": f"Unidad de aprendizaje {m.group(1)}:", "titulo": [], "logro": [], "temario": []}
            estado = "titulo"
            continue

        m = rx_unidad_inline.match(ln)
        if m:
            cerrar_unidad()
            curr = {"Unidad": f"Unidad de aprendizaje {m.group(1)}:", "titulo": [], "logro": [], "temario": []}
            push_unique(curr["titulo"], m.group(2))
            estado = "titulo"
            continue

        if curr is None:
            continue

        m = rx_logro_full.match(ln)
        if m:
            estado = "logro"
            continue
        m = rx_logro_inline.match(ln)
        if m:
            estado = "logro"
            push_unique(curr["logro"], m.group(1))
            continue

        m = rx_temario_full.match(ln)
        if m:
            estado = "temario"
            continue
        m = rx_temario_inline.match(ln)
        if m:
            estado = "temario"
            push_unique(curr["temario"], m.group(1))
            continue

        if estado == "titulo":
            push_unique(curr["titulo"], ln)
        elif estado == "logro":
            push_unique(curr["logro"], ln)
        elif estado == "temario":
            push_unique(curr["temario"], ln)
        else:
            push_unique(curr["titulo"], ln)

    cerrar_unidad()

    df = pd.DataFrame(
        registros,
        columns=["Archivo", "Unidad", "Unidades_Aprendizaje", "Logro_Especifico", "Temario"],
    )

    if not df.empty:
        last_idx = df.index[-1]
        tem_full = df.at[last_idx, "Temario"] or ""
        mcut = rx_cut_inside.search(tem_full)
        if mcut:
            df.at[last_idx, "Temario"] = norm_line(tem_full[:mcut.start()])

    if df.empty:
        sec_text = "\n".join(s5)
        rx_u = re.compile(r"Unidad\s*de\s*aprendizaje\s*(\d+)\s*:\s*", re.IGNORECASE)
        hits = list(rx_u.finditer(sec_text))
        filas = []
        for i, h in enumerate(hits):
            b_start = h.end()
            b_end = hits[i + 1].start() if i + 1 < len(hits) else len(sec_text)
            bloque = sec_text[b_start:b_end]

            m_log = re.search(
                r"Logros?\s*espec[íi]ficos?\s*de\s*aprendizaje\s*:\s*", bloque, re.IGNORECASE
            )
            m_tem = re.search(r"Temario\s*:\s*", bloque, re.IGNORECASE)
            cut_to = min([x.start() for x in [m_log, m_tem] if x] + [len(bloque)])

            titulo = limpia_semana(bloque[:cut_to].strip())
            titulo = norm_line(titulo)

            logro = ""
            if m_log:
                lstart = m_log.end()
                lend = m_tem.start() if m_tem else len(bloque)
                logro = collapse_full_repeats(norm_line(bloque[lstart:lend]))

            temario = ""
            if m_tem:
                tstart = m_tem.end()
                temario = collapse_full_repeats(norm_line(bloque[tstart:]))

            filas.append({
                "Archivo": nombre_archivo_base,
                "Unidad": f"Unidad de aprendizaje {h.group(1)}:",
                "Unidades_Aprendizaje": titulo,
                "Logro_Especifico": logro,
                "Temario": temario,
            })
        df = pd.DataFrame(
            filas,
            columns=["Archivo", "Unidad", "Unidades_Aprendizaje", "Logro_Especifico", "Temario"],
        )

    return df


def procesar_silabo_uploaded_file(uploaded_file) -> tuple[str, bytes]:
    """
    Procesa un archivo Word de sílabo (UploadedFile de Streamlit) y devuelve:
      - nombre de archivo Excel de salida
      - contenido binario del Excel (bytes)
    """
    nombre_archivo = uploaded_file.name
    nombre_base = os.path.splitext(nombre_archivo)[0]

    file_bytes = uploaded_file.getvalue()
    file_stream = BytesIO(file_bytes)
    doc = Document(file_stream)

    cronograma_raw = extraer_cronograma_doc(doc, nombre_archivo)
    df_cronograma = pd.DataFrame(cronograma_raw)

    if not df_cronograma.empty:
        columnas_orden = [
            "Cod_Catalogo", "Nombre_curso",
            "Unidad", "Titulo_unidad",
            "Semana", "Sesión", "Tema", "Actividades y evaluaciones"
        ]
        df_cronograma = df_cronograma[columnas_orden]

    df_evaluacion = extraer_tabla_evaluacion_doc(doc, nombre_archivo)
    df_datos_generales = extraer_datos_generales_doc(doc, nombre_archivo)
    df_logro_metodologia = extraer_logro_metodologia_doc(doc, nombre_archivo)
    df_pesos = extraer_pesos_doc(doc, nombre_archivo)
    df_unidades_logros = extraer_unidades_logros_doc(doc, nombre_archivo)

    cols_unidades = ["Archivo", "Unidad", "Unidades_Aprendizaje", "Logro_Especifico", "Temario"]
    if df_unidades_logros.empty:
        df_unidades_logros = pd.DataFrame(columns=cols_unidades)

    buffer = BytesIO()
    stamp = time.strftime("%Y%m%d_%H%M")
    nombre_salida = f"{nombre_base}_CRONOGRAMA_{stamp}.xlsx"

    try:
        writer = pd.ExcelWriter(buffer, engine="openpyxl")
    except Exception:
        writer = pd.ExcelWriter(buffer)

    with writer as w:
        if not df_cronograma.empty:
            df_cronograma.to_excel(w, sheet_name="Cronograma", index=False)
        if not df_evaluacion.empty:
            df_evaluacion.to_excel(w, sheet_name="Sistema_Eva", index=False)
        if not df_datos_generales.empty:
            df_datos_generales.to_excel(w, sheet_name="Datos_generales", index=False)
        if not df_logro_metodologia.empty:
            df_logro_metodologia.to_excel(w, sheet_name="Logro_Metodologia", index=False)
        if not df_pesos.empty:
            df_pesos.to_excel(w, sheet_name="Pesos", index=False)

        df_unidades_logros.to_excel(w, sheet_name="Unidades_Logros", index=False)

    buffer.seek(0)
    return nombre_salida, buffer.getvalue()


# ======================================================
# UTILIDADES – ESTADO MÓDULO SYLLABUS
# ======================================================

def reset_syllabus_module():
    """
    Limpia el estado asociado al módulo 'Convert Syllabus to Excel'
    y fuerza un rerun para que se vea como si se entrara por primera vez.
    """
    # Limpiar resultados y firma de archivos
    if "syllabus_results" in st.session_state:
        del st.session_state["syllabus_results"]

    st.session_state["syllabus_last_files_signature"] = None

    # Cambiar el key del file_uploader para que se vacíe
    if "syllabus_reset_counter" in st.session_state:
        st.session_state["syllabus_reset_counter"] += 1
    else:
        st.session_state["syllabus_reset_counter"] = 1

    try:
        st.rerun()
    except Exception:
        st.experimental_rerun()


def build_files_signature(uploaded_files) -> str:
    """
    Genera una firma simple de los archivos cargados para detectar cambios
    (basado en nombre y tamaño).
    """
    parts = []
    for uf in uploaded_files:
        size = getattr(uf, "size", None)
        parts.append(f"{uf.name}:{size}")
    parts.sort()
    return "|".join(parts)


def render_task_progress(placeholder, idx: int, total: int, filename: str):
    """
    Renderiza la barra de progreso tipo 'task' en un placeholder dado.
    """
    if total <= 0:
        pct_int = 0
    else:
        pct = idx / total
        pct_int = int(pct * 100)

    file_label = filename or ""
    html = f"""
    <div class="progress-bar-ui-task">
        <div class="progress-bar-ui-task-header">
            <span class="progress-bar-ui-task-title">Procesando sílabos...</span>
            <span class="progress-bar-ui-task-percentage">{pct_int}%</span>
        </div>
        <div class="progress-bar-ui-task-bar-track">
            <div class="progress-bar-ui-task-bar-fill" style="width:{pct_int}%;"></div>
        </div>
        <div class="progress-bar-ui-task-sub">
            Procesando archivo {idx}/{total} · <code>{file_label}</code>
        </div>
    </div>
    """
    placeholder.markdown(html, unsafe_allow_html=True)


# ======================================================
# UI – CONVERT SYLLABUS TO EXCEL
# ======================================================

def render_syllabus_to_excel():
    render_hero_syllabus()

    # -----------------------------
    # Botón REINICIAR debajo del hero (lado izquierdo)
    # -----------------------------
    cols_top = st.columns([1, 3])
    with cols_top[0]:
        if st.button("Reiniciar", key="btn_syllabus_reset"):
            reset_syllabus_module()
    with cols_top[1]:
        st.write("")

    # -----------------------------
    # Paso 1: Cargar Word(s) del sílabo
    # -----------------------------
    st.markdown('<div class="utp-step-card">', unsafe_allow_html=True)
    step1_ph = st.empty()

    reset_counter = st.session_state.get("syllabus_reset_counter", 0)

    uploaded_files = st.file_uploader(
        "Cargar archivo(s) Word del sílabo (.docx)",
        type=["docx"],
        accept_multiple_files=True,
        help="Puedes seleccionar uno o varios archivos de sílabos en formato Word.",
        key=f"file_uploader_syllabus_{reset_counter}",
    )

    have_files = bool(uploaded_files)
    step1_ph.markdown(
        render_step_header_html("1", "Cargar archivo(s) Word del sílabo (.docx)", have_files),
        unsafe_allow_html=True,
    )

    if have_files:
        st.success(f"{len(uploaded_files)} archivo(s) cargado(s).")
    else:
        st.info("Selecciona al menos un archivo de sílabo para continuar.")

    st.markdown("</div>", unsafe_allow_html=True)

    # -----------------------------
    # Paso 2 (lógica automática, SIN mostrar card propia)
    # Se ejecuta en cuanto haya archivos cargados.
    # -----------------------------
    resultados = st.session_state.get("syllabus_results", None)

    if have_files:
        signature = build_files_signature(uploaded_files)
        last_signature = st.session_state.get("syllabus_last_files_signature", None)

        need_processing = (resultados is None) or (signature != last_signature)

        if need_processing:
            resultados = []
            progress_placeholder = st.empty()

            total = len(uploaded_files)
            with st.spinner("Procesando sílabos y generando archivos Excel..."):
                for idx, uf in enumerate(uploaded_files, start=1):
                    render_task_progress(progress_placeholder, idx, total, uf.name)
                    try:
                        nombre_excel, contenido_excel = procesar_silabo_uploaded_file(uf)
                        resultados.append(
                            {
                                "input_name": uf.name,
                                "output_name": nombre_excel,
                                "content": contenido_excel,
                                "error": None,
                            }
                        )
                    except Exception as e:
                        resultados.append(
                            {
                                "input_name": uf.name,
                                "output_name": None,
                                "content": None,
                                "error": str(e),
                            }
                        )

            progress_placeholder.empty()

            st.session_state.syllabus_results = resultados
            st.session_state.syllabus_last_files_signature = signature
            st.success("Transformación finalizada. Revisa la sección de descarga de Excel.")
    else:
        st.session_state.syllabus_last_files_signature = None

    # -----------------------------
    # Paso 2 (visible): Descarga de Excel
    # -----------------------------
    st.markdown('<div class="utp-step-card">', unsafe_allow_html=True)
    render_simple_step_header("2", "Descargar archivos Excel")

    resultados = st.session_state.get("syllabus_results", None)

    if not resultados:
        st.info("Aún no hay archivos generados. Carga los sílabos para iniciar la transformación automática.")
    else:
        exitosos = [(r.get("output_name"), r.get("content")) for r in resultados
                    if r.get("output_name") and r.get("content")]
        errores = [r for r in resultados if not (r.get("output_name") and r.get("content"))]

        if exitosos:
            zip_buffer = BytesIO()
            with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zf:
                for fname, fcontent in exitosos:
                    zf.writestr(fname, fcontent)
            zip_buffer.seek(0)

            stamp = time.strftime("%Y%m%d_%H%M")
            zip_name = f"SYLLABUS_EXCEL_{stamp}.zip"

            st.download_button(
                label="Download Excel (ZIP)",
                data=zip_buffer.getvalue(),
                file_name=zip_name,
                mime="application/zip",
                key=f"download_zip_{zip_name}",
            )
        else:
            st.error("No se pudo generar ningún archivo Excel válido a partir de los sílabos cargados.")

        if errores:
            with st.expander("Ver detalles de errores"):
                for r in errores:
                    st.write(
                        f"• Archivo original: `{r.get('input_name', '(desconocido)')}` – "
                        f"{r.get('error') or 'Error desconocido.'}"
                    )

    st.markdown("</div>", unsafe_allow_html=True)

# ======================================================
# MAIN
# ======================================================

def main():
    st.set_page_config(
        page_title="Plataforma UTP – Syllabus to Excel",
        page_icon="📄",
        layout="wide",
        initial_sidebar_state="expanded",
    )

    apply_global_styles()

    if "syllabus_results" not in st.session_state:
        st.session_state.syllabus_results = None
    if "syllabus_reset_counter" not in st.session_state:
        st.session_state.syllabus_reset_counter = 0
    if "syllabus_last_files_signature" not in st.session_state:
        st.session_state.syllabus_last_files_signature = None

    with st.sidebar:
        header_placeholder = st.empty()

        page = st.radio(
            "Módulos",
            [
                "Home",
                "Convert Syllabus to Excel",
            ],
        )

        if page == "Convert Syllabus to Excel":
            header_html = get_sidebar_header_html(
                title="Syllabus to Excel Transformation",
                subtitle="Transformacion de silabos a tablas Excel.",
                icon="📊",
            )
        else:
            header_html = get_sidebar_header_html(
                title="Plataforma UTP",
                subtitle="Transformación de silabos a tablas Excel",
                icon="📚",
            )
        header_placeholder.markdown(header_html, unsafe_allow_html=True)

    # Contenido principal según módulo
    if page == "Home":
        render_home()
    elif page == "Convert Syllabus to Excel":
        render_syllabus_to_excel()

if __name__ == "__main__":
    main()













